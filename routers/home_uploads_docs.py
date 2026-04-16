"""Document Studio — read, edit, combine, sign, and convert uploaded files.

Companion router for home_uploads.py.  All routes live under /home/uploads/{page_id}.
Kept separate from home_uploads.py (already 8.9 KB) for cohesion and line-limit hygiene.

Phase 4 endpoints (all implemented here):
  GET  /{pid}/files/{src}/{id}/content    read text or Word doc as text/HTML
  PUT  /{pid}/files/page/{id}/content     save edited text back to disk
  POST /{pid}/files/page/combine          merge PDFs or join text files
  POST /{pid}/files/page/{id}/convert     docx↔pdf, txt→pdf, pdf→txt
  POST /{pid}/files/page/{id}/sign        stamp a drawn signature onto a PDF page

Phase 6 endpoint (Feature A):
  GET  /{pid}/files/page/{id}/wopi-token  issue short-lived WOPI token for Collabora
"""
from __future__ import annotations

import base64
import csv
import io
import re
import shutil
import uuid
from html import escape
from html.parser import HTMLParser

from fastapi import APIRouter, HTTPException, Request
from fastapi.responses import JSONResponse, Response
from pydantic import BaseModel

from routers.attachments_db import UPLOAD_DIR
from routers.home_uploads import _demo_guard, _require_uploads_page, _valid_src
from routers.uploads_db import (
    create_page_upload,
    get_note_attachment_owned,
    get_page_upload_owned,
)
from routers.uploads_docs_db import get_page_upload_owned_bulk, update_page_upload_size

router = APIRouter(prefix="/home/uploads", tags=["uploads-docs"])

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MAX_EDIT_BYTES = 1_000_000  # 1 MB guard on textarea saves


# ── Pydantic models ───────────────────────────────────────────────────────────

class ContentBody(BaseModel):
    content: str


class SpreadsheetBody(BaseModel):
    content_b64: str   # base64-encoded file bytes (SheetJS XLSX or UTF-8 CSV)
    format: str        # "xlsx" | "csv"


MAX_SPREADSHEET_BYTES = 10_000_000  # 10 MB guard (larger than the 1 MB text guard)


class CombineBody(BaseModel):
    ids: list[int]
    output_name: str = ""
    combine_type: str  # "pdf" | "text"


class ConvertBody(BaseModel):
    to_format: str  # "pdf" | "txt"


class SignPlacement(BaseModel):
    x_pct: float   # 0-1 from left edge
    y_pct: float   # 0-1 from top edge (CSS origin; flipped to PDF bottom-up)
    page_num: int = 0


class SaveAsTxtBody(BaseModel):
    content: str


class SignBody(BaseModel):
    signature_data: str              # data:image/png;base64,…
    page_num: int = 0                # legacy single-placement fallback
    x_pct: float = 0.65             # legacy single-placement fallback
    y_pct: float = 0.80             # legacy single-placement fallback
    placements: list[SignPlacement] | None = None  # multi-placement (preferred)


# ── GET /{pid}/files/{src}/{id}/content ───────────────────────────────────────

@router.get("/{page_id}/files/{src}/{file_id}/content")
async def read_content(request: Request, page_id: int, src: str, file_id: int):
    """Return full file content as plain text or docx-derived HTML."""
    uid = request.session.get("user_id")
    if not uid:
        raise HTTPException(status_code=401)
    await _require_uploads_page(page_id, uid)
    src = _valid_src(src)

    row = await (
        get_note_attachment_owned(file_id, uid)
        if src == "note"
        else get_page_upload_owned(file_id, uid)
    )
    if not row:
        raise HTTPException(status_code=404)

    mime = row["mime_type"]
    disk_path = UPLOAD_DIR / row["filename"]

    if mime.startswith("text/") or mime == "application/json":
        try:
            text = disk_path.read_text(encoding="utf-8", errors="replace")
        except Exception as exc:
            raise HTTPException(status_code=500, detail="Could not read file") from exc
        return JSONResponse({"content": text, "content_type": "text"})

    if mime == _DOCX_MIME:
        try:
            from docx import Document
            doc = Document(disk_path)
            html = _docx_body_to_html(doc)
        except Exception as exc:
            raise HTTPException(status_code=500, detail="Could not parse Word document") from exc
        return JSONResponse({"content": html, "content_type": "docx_html"})

    if mime == "text/csv" or (mime.startswith("text/") and disk_path.suffix.lower() == ".csv"):
        try:
            raw = disk_path.read_text(encoding="utf-8-sig", errors="replace")
            html = _csv_to_html(raw)
        except Exception as exc:
            raise HTTPException(status_code=500, detail="Could not parse CSV") from exc
        return JSONResponse({"content": html, "content_type": "csv_html"})

    raise HTTPException(status_code=400, detail="Content read not supported for this file type")


def _docx_body_to_html(doc) -> str:  # type: ignore[annotation-unchecked]
    """Pure-XML walk of the body — no python-docx proxy constructors, no parent-chain issues."""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def _w(name: str) -> str:
        return f"{{{W}}}{name}"

    def _local(el) -> str:
        return el.tag.split("}")[-1] if "}" in el.tag else el.tag

    # Build style-id → display-name from the real style registry (safe, no proxy walk needed)
    _HEADING_TAGS = {
        "Heading 1": "h2", "Heading 2": "h3", "Heading 3": "h4",
        "Heading 4": "h5", "Heading 5": "h6", "Heading 6": "h6",
        "Title": "h1",    "Subtitle": "h2",
    }
    sid_to_name: dict[str, str] = {s.style_id: s.name for s in doc.styles}

    def _drawing_html(drawing_el) -> str:
        """Extract embedded image from a w:drawing element and return an <img> data-URI tag."""
        A_NS      = "http://schemas.openxmlformats.org/drawingml/2006/main"
        R_NS      = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        blip_tag  = f"{{{A_NS}}}blip"
        embed_key = f"{{{R_NS}}}embed"
        for child in drawing_el.iter():
            if child.tag != blip_tag:
                continue
            r_id = child.get(embed_key)
            if not r_id:
                break
            try:
                part = doc.part.related_parts[r_id]
                blob = part.blob
                if len(blob) > 2_000_000:   # skip images > 2 MB in the viewer
                    return ""
                img_b64 = base64.b64encode(blob).decode()
                mime    = part.content_type or "image/png"
                return (
                    f'<img src="data:{mime};base64,{img_b64}"'
                    ' style="max-width:100%;height:auto;display:block;margin:0.5em 0;"'
                    ' alt="embedded image">'
                )
            except Exception:
                return ""
        return ""

    def _run_html(r_el) -> str:
        html_parts: list[str] = []
        text_parts: list[str] = []
        for el in r_el:
            loc = _local(el)
            if loc == "t":
                text_parts.append(escape(el.text or ""))
            elif loc == "br":
                text_parts.append("<br>")
            elif loc == "drawing":
                img = _drawing_html(el)
                if img:
                    html_parts.append(img)
        text = "".join(text_parts)
        if text:
            rpr = r_el.find(_w("rPr"))
            if rpr is not None:
                bold   = rpr.find(_w("b"))  is not None
                italic = rpr.find(_w("i"))  is not None
                under  = rpr.find(_w("u"))  is not None
                if bold and italic:
                    text = f"<strong><em>{text}</em></strong>"
                elif bold:
                    text = f"<strong>{text}</strong>"
                elif italic:
                    text = f"<em>{text}</em>"
                if under:
                    text = f"<u>{text}</u>"
            html_parts.append(text)
        return "".join(html_parts)

    def _para_html(p_el) -> str:
        ppr = p_el.find(_w("pPr"))
        style_name = ""
        if ppr is not None:
            ps = ppr.find(_w("pStyle"))
            if ps is not None:
                sid = ps.get(_w("val")) or ""
                style_name = sid_to_name.get(sid, sid)
        tag   = _HEADING_TAGS.get(style_name, "p")
        inner = "".join(_run_html(r) for r in p_el.findall(_w("r"))) or "&nbsp;"
        return f"<{tag}>{inner}</{tag}>"

    def _cell_html(tc_el, is_header: bool) -> str:
        tag   = "th" if is_header else "td"
        texts: list[str] = []
        for p in tc_el.findall(_w("p")):
            for r in p.findall(_w("r")):
                t = r.find(_w("t"))
                if t is not None:
                    texts.append(t.text or "")
        return f"<{tag}>{escape(' '.join(texts))}</{tag}>"

    def _table_html(tbl_el) -> str:
        rows_html: list[str] = []
        for i, tr in enumerate(tbl_el.findall(_w("tr"))):
            cells = [_cell_html(tc, i == 0) for tc in tr.findall(_w("tc"))]
            rows_html.append("<tr>" + "".join(cells) + "</tr>")
        return '<table class="bw-doc-table">' + "".join(rows_html) + "</table>"

    parts: list[str] = []
    for child in doc.element.body:
        loc = _local(child)
        if loc == "p":
            parts.append(_para_html(child))
        elif loc == "tbl":
            parts.append(_table_html(child))
    return "".join(parts) or "<p><em>Empty document</em></p>"


def _csv_to_html(raw: str) -> str:
    """Parse CSV text → styled HTML table (max 500 rows for safety)."""
    reader = list(csv.reader(io.StringIO(raw)))
    if not reader:
        return "<p><em>Empty CSV</em></p>"
    MAX_ROWS = 500
    header, rows = reader[0], reader[1:MAX_ROWS + 1]
    truncated   = len(reader) - 1 > MAX_ROWS
    th = "".join(f"<th>{escape(c)}</th>" for c in header)
    body_rows   = []
    for row in rows:
        # Pad short rows to header length
        padded = row + [""] * max(0, len(header) - len(row))
        body_rows.append("<tr>" + "".join(f"<td>{escape(c)}</td>" for c in padded) + "</tr>")
    note = f'<p class="bw-csv-note">Showing first {MAX_ROWS} of {len(reader)-1} rows.</p>' if truncated else ""
    return f'<table class="bw-doc-table"><thead><tr>{th}</tr></thead><tbody>' + "".join(body_rows) + f"</tbody></table>{note}"


# ── PUT /{pid}/files/page/{id}/content ────────────────────────────────────────

@router.put("/{page_id}/files/page/{file_id}/content")
async def save_content(request: Request, page_id: int, file_id: int, body: ContentBody):
    """Overwrite a text file's content. page-src only."""
    if guard := _demo_guard(request):
        return guard
    uid = request.session.get("user_id")
    if not uid:
        raise HTTPException(status_code=401)
    await _require_uploads_page(page_id, uid)

    row = await get_page_upload_owned(file_id, uid)
    if not row:
        raise HTTPException(status_code=404)

    mime = row["mime_type"]
    if not (mime.startswith("text/") or mime == "application/json"):
        raise HTTPException(status_code=400, detail="Only text files are editable")

    data = body.content.encode("utf-8")
    if len(data) > MAX_EDIT_BYTES:
        raise HTTPException(status_code=413, detail="Content exceeds the 1 MB edit limit")

    try:
        (UPLOAD_DIR / row["filename"]).write_bytes(data)
    except Exception as exc:
        raise HTTPException(status_code=500, detail="Could not write file") from exc

    await update_page_upload_size(file_id, uid, len(data))
    return JSONResponse({"ok": True, "size": len(data)})


# ── PUT /{pid}/files/page/{id}/spreadsheet ────────────────────────────────────

_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

@router.put("/{page_id}/files/page/{file_id}/spreadsheet")
async def save_spreadsheet(
    request: Request, page_id: int, file_id: int, body: SpreadsheetBody
):
    """Receive base64-encoded XLSX or CSV bytes and overwrite the file on disk."""
    if guard := _demo_guard(request):
        return guard
    uid = request.session.get("user_id")
    if not uid:
        raise HTTPException(status_code=401)
    await _require_uploads_page(page_id, uid)

    row = await get_page_upload_owned(file_id, uid)
    if not row:
        raise HTTPException(status_code=404)

    mime = row["mime_type"]
    if mime not in (_XLSX_MIME, "text/csv"):
        raise HTTPException(status_code=400, detail="Only XLSX and CSV files can be saved via this endpoint")

    if body.format not in ("xlsx", "csv"):
        raise HTTPException(status_code=400, detail="format must be 'xlsx' or 'csv'")

    try:
        data = base64.b64decode(body.content_b64)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid base64 content")

    if len(data) > MAX_SPREADSHEET_BYTES:
        raise HTTPException(status_code=413, detail="Spreadsheet exceeds the 10 MB limit")

    try:
        (UPLOAD_DIR / row["filename"]).write_bytes(data)
    except Exception as exc:
        raise HTTPException(status_code=500, detail="Could not write file") from exc

    await update_page_upload_size(file_id, uid, len(data))
    return JSONResponse({"ok": True, "size": len(data)})


# ── POST /{pid}/files/page/combine ────────────────────────────────────────────

@router.post("/{page_id}/files/page/combine")
async def combine_files(request: Request, page_id: int, body: CombineBody):
    """Merge 2–20 PDFs into one, or concatenate 2–20 text files."""
    if guard := _demo_guard(request):
        return guard
    uid = request.session.get("user_id")
    if not uid:
        raise HTTPException(status_code=401)
    await _require_uploads_page(page_id, uid)

    if not 2 <= len(body.ids) <= 20:
        raise HTTPException(status_code=400, detail="Select 2–20 files to combine")

    rows = await get_page_upload_owned_bulk(body.ids, uid)
    if len(rows) != len(body.ids):
        raise HTTPException(status_code=404, detail="One or more files not found")

    out_stem = (body.output_name.strip() or "combined")[:80]

    if body.combine_type == "pdf":
        if any(r["mime_type"] != "application/pdf" for r in rows):
            raise HTTPException(status_code=400, detail="All selected files must be PDFs for PDF merge")
        import pypdf
        writer = pypdf.PdfWriter()
        for r in rows:
            reader = pypdf.PdfReader(str(UPLOAD_DIR / r["filename"]))
            for pg in reader.pages:
                writer.add_page(pg)
        buf = io.BytesIO()
        writer.write(buf)
        data = buf.getvalue()
        stored = f"{uuid.uuid4().hex}.pdf"
        (UPLOAD_DIR / stored).write_bytes(data)
        new_id = await create_page_upload(
            page_id, uid, stored, f"{out_stem}.pdf", "application/pdf", len(data),
        )
        return JSONResponse({"ok": True, "file": {
            "id": new_id, "original_name": f"{out_stem}.pdf", "size": len(data),
        }})

    if body.combine_type == "text":
        if any(not r["mime_type"].startswith("text/") for r in rows):
            raise HTTPException(status_code=400, detail="All selected files must be text files for text join")
        parts: list[str] = []
        for r in rows:
            parts.append(f"── {r['original_name']} ──\n")
            parts.append((UPLOAD_DIR / r["filename"]).read_text(encoding="utf-8", errors="replace"))
            parts.append("\n\n")
        data = "".join(parts).encode("utf-8")
        stored = f"{uuid.uuid4().hex}.txt"
        (UPLOAD_DIR / stored).write_bytes(data)
        new_id = await create_page_upload(
            page_id, uid, stored, f"{out_stem}.txt", "text/plain", len(data),
        )
        return JSONResponse({"ok": True, "file": {
            "id": new_id, "original_name": f"{out_stem}.txt", "size": len(data),
        }})

    raise HTTPException(status_code=400, detail="combine_type must be 'pdf' or 'text'")


# ── POST /{pid}/files/page/{id}/convert ───────────────────────────────────────

@router.post("/{page_id}/files/page/{file_id}/convert")
async def convert_file(request: Request, page_id: int, file_id: int, body: ConvertBody):
    """Convert docx/txt → pdf, or docx/pdf → txt.  Always creates a new file."""
    if guard := _demo_guard(request):
        return guard
    uid = request.session.get("user_id")
    if not uid:
        raise HTTPException(status_code=401)
    await _require_uploads_page(page_id, uid)

    row = await get_page_upload_owned(file_id, uid)
    if not row:
        raise HTTPException(status_code=404)

    mime = row["mime_type"]
    disk_path = UPLOAD_DIR / row["filename"]
    src_stem = row["original_name"].rsplit(".", 1)[0]
    to_fmt = body.to_format.lower()

    if to_fmt not in ("pdf", "txt"):
        raise HTTPException(status_code=400, detail="to_format must be 'pdf' or 'txt'")

    if to_fmt == "txt":
        if mime == _DOCX_MIME:
            from docx import Document
            doc = Document(disk_path)
            text = "\n".join(p.text for p in doc.paragraphs)
        elif mime == "application/pdf":
            import pypdf
            reader = pypdf.PdfReader(str(disk_path))
            text = "\n\n".join(pg.extract_text() or "" for pg in reader.pages)
        else:
            raise HTTPException(status_code=400, detail="Only PDF or DOCX can be converted to TXT")
        data = text.encode("utf-8")
        stored = f"{uuid.uuid4().hex}.txt"
        (UPLOAD_DIR / stored).write_bytes(data)
        new_id = await create_page_upload(
            page_id, uid, stored, f"{src_stem}.txt", "text/plain", len(data),
        )
        return JSONResponse({"ok": True, "file": {
            "id": new_id, "original_name": f"{src_stem}.txt", "size": len(data),
        }})

    # to_fmt == "pdf"
    if mime.startswith("text/"):
        text = disk_path.read_text(encoding="utf-8", errors="replace")
    elif mime == _DOCX_MIME:
        from docx import Document
        doc = Document(disk_path)
        text = "\n".join(p.text for p in doc.paragraphs)
    else:
        raise HTTPException(status_code=400, detail="Only TXT or DOCX can be converted to PDF")

    data = _text_to_pdf_bytes(text)
    stored = f"{uuid.uuid4().hex}.pdf"
    (UPLOAD_DIR / stored).write_bytes(data)
    new_id = await create_page_upload(
        page_id, uid, stored, f"{src_stem}.pdf", "application/pdf", len(data),
    )
    return JSONResponse({"ok": True, "file": {
        "id": new_id, "original_name": f"{src_stem}.pdf", "size": len(data),
    }})


# ── PUT /{pid}/files/page/{id}/docx-content — save edited HTML back to DOCX ──

class DocxHtmlBody(BaseModel):
    html: str   # full innerHTML of the editor's contenteditable div


@router.put("/{page_id}/files/page/{file_id}/docx-content")
async def save_docx_content(
    request: Request, page_id: int, file_id: int, body: DocxHtmlBody
):
    """Convert caller-supplied HTML back to DOCX and overwrite the original file."""
    if guard := _demo_guard(request):
        return guard
    uid = request.session.get("user_id")
    if not uid:
        raise HTTPException(status_code=401)
    await _require_uploads_page(page_id, uid)

    row = await get_page_upload_owned(file_id, uid)
    if not row or row["mime_type"] != _DOCX_MIME:
        raise HTTPException(status_code=400, detail="File must be a Word document")

    if len(body.html.encode()) > 2_000_000:
        raise HTTPException(status_code=413, detail="Content too large (max 2 MB)")

    try:
        data = _html_to_docx_bytes(body.html)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"HTML→DOCX conversion failed: {exc}") from exc

    try:
        (UPLOAD_DIR / row["filename"]).write_bytes(data)
    except Exception as exc:
        raise HTTPException(status_code=500, detail="Could not write file") from exc

    await update_page_upload_size(file_id, uid, len(data))
    return JSONResponse({"ok": True, "size": len(data)})


# ── HTML → DOCX converter — used by the Word-like editor ────────────────────

def _css_color_to_docx(val: str):
    """Parse CSS colour (rgb/hex) → docx RGBColor or None."""
    from docx.shared import RGBColor
    val = val.strip()
    m = re.match(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', val)
    if m:
        return RGBColor(int(m.group(1)), int(m.group(2)), int(m.group(3)))
    m = re.match(r'#([0-9a-fA-F]{6})', val)
    if m:
        h = m.group(1)
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    return None


def _html_to_docx_bytes(html: str) -> bytes:
    """Parse HTML from the Word-like editor and emit a python-docx Document."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _HEADING_MAP = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}
    _ALIGN_MAP = {
        "left":    WD_ALIGN_PARAGRAPH.LEFT,
        "center":  WD_ALIGN_PARAGRAPH.CENTER,
        "right":   WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    doc     = Document()
    para    = doc.add_paragraph()
    level   = None   # None = Normal; 1-6 = heading
    li_type: list[str] = []  # stack of 'ul'|'ol'
    li_idx:  list[int] = []  # counters for ordered lists

    # Inline formatting state
    bold = italic = under = strike = False
    fg_color = bg_color = None
    font_size = None
    alignment = None

    def _flush_para():
        nonlocal para, alignment
        if alignment and para.runs:
            para.paragraph_format.alignment = alignment

    def _new_block(tag: str, attrs: dict):
        nonlocal para, level, alignment, bold, italic, under, strike, fg_color, bg_color, font_size
        _flush_para()
        alignment = None
        # extract text-align from style
        style_str = attrs.get("style", "")
        m = re.search(r'text-align:\s*(\w+)', style_str)
        if m:
            alignment = _ALIGN_MAP.get(m.group(1))
        if tag in _HEADING_MAP:
            level = _HEADING_MAP[tag]
            para  = doc.add_heading(level=level)
        else:
            level = None
            para  = doc.add_paragraph()

    class _Parser(HTMLParser):
        def handle_starttag(self, tag, attrs):
            nonlocal bold, italic, under, strike, fg_color, bg_color, font_size, alignment
            nonlocal para, level, li_type, li_idx
            attrs = dict(attrs)
            style = attrs.get("style", "")
            if tag in ("p", "div", *_HEADING_MAP):
                _new_block(tag, attrs)
            elif tag in ("ul", "ol"):
                li_type.append(tag)
                li_idx.append(0)
            elif tag == "li":
                _flush_para()
                indent = len(li_type)
                is_ol  = li_type[-1] == "ol" if li_type else False
                if is_ol:
                    li_idx[-1] += 1
                    prefix = f"{li_idx[-1]}. "
                else:
                    prefix = "\u2022 "
                para  = doc.add_paragraph(style="List Bullet" if not is_ol else "List Number")
                para.paragraph_format.left_indent = Pt(18 * indent)
            elif tag in ("strong", "b"):  bold   = True
            elif tag in ("em", "i"):      italic = True
            elif tag == "u":              under  = True
            elif tag in ("s", "del", "strike"): strike = True
            elif tag == "br":             para.add_run("\n")
            elif tag == "span":
                m = re.search(r'color:\s*([^;]+)', style)
                if m and "background" not in style[:m.start()]:
                    fg_color = _css_color_to_docx(m.group(1))
                m = re.search(r'background-color:\s*([^;]+)', style)
                if m:
                    bg_color = _css_color_to_docx(m.group(1))
                m = re.search(r'font-size:\s*([\d.]+)pt', style)
                if m:
                    font_size = float(m.group(1))

        def handle_endtag(self, tag):
            nonlocal bold, italic, under, strike, fg_color, bg_color, font_size
            nonlocal li_type, li_idx
            if tag in ("strong", "b"):  bold   = False
            elif tag in ("em", "i"):    italic = False
            elif tag == "u":            under  = False
            elif tag in ("s", "del", "strike"): strike = False
            elif tag == "span":         fg_color = bg_color = font_size = None
            elif tag in ("ul", "ol") and li_type:
                li_type.pop(); li_idx.pop()

        def handle_data(self, data):
            text = data  # preserve whitespace as-is from editor
            if not text:
                return
            run = para.add_run(text)
            run.bold      = bold
            run.italic    = italic
            run.underline = under
            if strike:
                run.font.strike = True
            if fg_color:
                run.font.color.rgb = fg_color
            if bg_color:
                run.font.highlight_color = None  # docx highlight is enum-only; skip bg
            if font_size:
                run.font.size = Pt(font_size)

        def handle_entityref(self, name):
            entities = {"nbsp": " ", "amp": "&", "lt": "<", "gt": ">",
                        "quot": '"', "apos": "'"}
            self.handle_data(entities.get(name, ""))

        def handle_charref(self, name):
            try:
                ch = chr(int(name[1:], 16) if name.startswith("x") else int(name))
            except Exception:
                ch = ""
            self.handle_data(ch)

    _Parser().feed(html)
    _flush_para()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def _text_to_pdf_bytes(text: str) -> bytes:
    """Render plain text to PDF bytes using reportlab.  Lazy import."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import pt
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=72, rightMargin=72, topMargin=72, bottomMargin=72,
    )
    styles = getSampleStyleSheet()
    story = []
    for line in text.splitlines():
        story.append(Paragraph(escape(line) if line.strip() else "&nbsp;", styles["Normal"]))
        story.append(Spacer(1, 4 * pt))
    if not story:
        story.append(Paragraph("(empty)", styles["Normal"]))
    doc.build(story)
    return buf.getvalue()


# ── POST /{pid}/files/page/{id}/sign ─────────────────────────────────────────

# ── GET page dimensions (used by placement UI to build a proportioned click-target) ──

@router.get("/{page_id}/files/page/{file_id}/page-dims")
async def get_page_dims(request: Request, page_id: int, file_id: int):
    """Return page-0 dimensions in points so the JS can render a correctly-proportioned preview."""
    uid = request.session.get("user_id")
    if not uid:
        raise HTTPException(status_code=401)
    await _require_uploads_page(page_id, uid)
    row = await get_page_upload_owned(file_id, uid)
    if not row or row["mime_type"] != "application/pdf":
        raise HTTPException(status_code=400, detail="File must be a PDF")
    try:
        import pypdf
        reader = pypdf.PdfReader(str(UPLOAD_DIR / row["filename"]))
        pg = reader.pages[0]
        return JSONResponse({
            "width_pt": float(pg.mediabox.width),
            "height_pt": float(pg.mediabox.height),
            "page_count": len(reader.pages),
        })
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


# ── GET/DELETE sign — backup status + stamp removal ───────────────────────────

@router.get("/{page_id}/files/page/{file_id}/sign")
async def sign_status(request: Request, page_id: int, file_id: int):
    """Return whether an unsigned backup exists for this PDF."""
    uid = request.session.get("user_id")
    if not uid:
        raise HTTPException(status_code=401)
    await _require_uploads_page(page_id, uid)
    row = await get_page_upload_owned(file_id, uid)
    if not row:
        raise HTTPException(status_code=404)
    bak = UPLOAD_DIR / (row["filename"] + ".bak")
    return JSONResponse({"has_backup": bak.exists()})


@router.delete("/{page_id}/files/page/{file_id}/sign")
async def remove_stamp(request: Request, page_id: int, file_id: int):
    """Restore the pre-signature backup, deleting the stamped version."""
    if guard := _demo_guard(request):
        return guard
    uid = request.session.get("user_id")
    if not uid:
        raise HTTPException(status_code=401)
    await _require_uploads_page(page_id, uid)
    row = await get_page_upload_owned(file_id, uid)
    if not row:
        raise HTTPException(status_code=404)
    bak = UPLOAD_DIR / (row["filename"] + ".bak")
    if not bak.exists():
        raise HTTPException(status_code=400, detail="No unsigned backup found")
    pdf_path = UPLOAD_DIR / row["filename"]
    shutil.copy2(str(bak), str(pdf_path))
    bak.unlink()
    size = pdf_path.stat().st_size
    await update_page_upload_size(file_id, uid, size)
    return JSONResponse({"ok": True, "size": size})


# ── POST /{pid}/files/page/{id}/save-as-txt ─────────────────────────────────

@router.post("/{page_id}/files/page/{file_id}/save-as-txt")
async def save_as_txt(request: Request, page_id: int, file_id: int, body: SaveAsTxtBody):
    """Save caller-supplied text as a brand-new .txt page upload.
    Used by the viewer's 'Edit as TXT' flow for Word documents.
    """
    if guard := _demo_guard(request):
        return guard
    uid = request.session.get("user_id")
    if not uid:
        raise HTTPException(status_code=401)
    await _require_uploads_page(page_id, uid)

    row = await get_page_upload_owned(file_id, uid)
    if not row:
        raise HTTPException(status_code=404)

    if len(body.content.encode()) > MAX_EDIT_BYTES:
        raise HTTPException(status_code=400, detail="Content too large (max 1 MB)")

    data      = body.content.encode("utf-8")
    stem      = row["original_name"].rsplit(".", 1)[0]
    out_name  = f"{stem}.txt"
    stored    = f"{uuid.uuid4().hex}.txt"
    (UPLOAD_DIR / stored).write_bytes(data)
    new_id = await create_page_upload(
        page_id, uid, stored, out_name, "text/plain", len(data)
    )
    return JSONResponse({"ok": True, "file": {
        "id": new_id, "original_name": out_name, "size": len(data),
    }})


# ── PDF signing helpers ───────────────────────────────────────────────────────

def _page_rotation(page) -> int:
    """Return the page's /Rotate value normalised to 0, 90, 180, or 270."""
    try:
        rot = getattr(page, "rotation", None)
        if rot is None:
            rot = int(page.get("/Rotate", 0) or 0)
        return int(rot) % 360
    except Exception:
        return 0


def _stamp_one_page(
    sig_img,           # PIL.Image already loaded (RGBA)
    target_page,       # pypdf page object
    x_pct: float,      # click x as fraction of visual width  (0-1, left→right)
    y_pct: float,      # click y as fraction of visual height (0-1, top→bottom)
):
    """
    Build a reportlab overlay for ONE page, handling page rotation so the
    signature always appears right-side-up at the visual click position.

    Coordinate transforms (derived from PDF /Rotate CW rotation matrices):
      R=0:   raw_cx = W*x,      raw_cy = H*(1-y)
      R=90:  raw_cx = W*y,      raw_cy = H*x        sig pre-rotated CCW 90°
      R=180: raw_cx = W*(1-x),  raw_cy = H*y        sig pre-rotated 180°
      R=270: raw_cx = W*(1-y),  raw_cy = H*(1-x)   sig pre-rotated CW 90°
    """
    import PIL.Image
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas as rl_canvas
    import io as _io

    rot    = _page_rotation(target_page)
    pg_w   = float(target_page.mediabox.width)
    pg_h   = float(target_page.mediabox.height)

    # The PDF viewer applies /Rotate CW when displaying the page (PDF spec §14.5).
    # To keep the signature upright in the viewer we must pre-rotate it by the
    # SAME amount COUNTER-CLOCKWISE in raw space, so the viewer's CW rotation
    # cancels it out.
    # PIL positive angles = CCW, so pil_rot = rot directly:
    #   R=90  → PIL(90)  CCW in raw → viewer +90° CW → 0° (upright) ✓
    #   R=180 → PIL(180)                                              ✓
    #   R=270 → PIL(270) CCW in raw → viewer +270° CW → 0° (upright) ✓
    pil_rot = rot
    if pil_rot:
        draw_img = sig_img.rotate(pil_rot, expand=True)
    else:
        draw_img = sig_img

    # Target sig width = 25% of the *visual* page width.
    #   For R=90/270 the display swaps (vis_w = pg_h, vis_h = pg_w).
    vis_w = pg_h if rot in (90, 270) else pg_w
    vis_sig_w = vis_w * 0.25
    # Keep the signature's original proportions regardless of PIL rotation.
    # Using draw_img here would give the INVERTED aspect for R=90/270 because
    # PIL.rotate(90/270, expand=True) swaps image width/height.
    orig_aspect = sig_img.width / sig_img.height  # always from the unrotated canvas
    vis_sig_h = vis_sig_w / orig_aspect

    # Convert visual sig size to raw PDF pts (axis swap for 90°/270°).
    if rot in (90, 270):
        raw_sig_w = vis_sig_h   # raw x ↔ visual y
        raw_sig_h = vis_sig_w   # raw y ↔ visual x
    else:
        raw_sig_w = vis_sig_w
        raw_sig_h = vis_sig_h

    # Coordinate transforms (visual CSS % → raw PDF pts).
    # Derived from PDF /Rotate spec rotation matrices:
    #   R=0:   standard, just flip y for PDF bottom-up origin
    #   R=90 CW:  axes swap, no inversion
    #   R=180:  both axes inverted, but y inverted back by PDF flip
    #   R=270 CW: axes swap + both inverted
    if rot == 0:
        raw_cx = pg_w * x_pct
        raw_cy = pg_h * (1.0 - y_pct)
    elif rot == 90:
        # R=90 CW: viewer maps raw→screen as: screen_x=ry_pdf, screen_y=rx
        # Inverse (x_pct=screen_x/vis_w, y_pct=screen_y/vis_h, vis_w=H, vis_h=W):
        #   ry_pdf = x_pct * H  →  raw_cy = pg_h * x_pct
        #   rx     = y_pct * W  →  raw_cx = pg_w * y_pct
        raw_cx = pg_w * y_pct
        raw_cy = pg_h * x_pct
    elif rot == 180:
        # CCW 180° display: both axes flip
        raw_cx = pg_w * (1.0 - x_pct)
        raw_cy = pg_h * y_pct           # correct (y flipped twice = same)
    else:  # 270
        # R=270 CW (= 90 CCW): viewer maps raw→screen as:
        #   screen_x = H - ry_pdf, screen_y = W - rx
        # Inverse (vis_w=H, vis_h=W):
        #   ry_pdf = H*(1-x_pct)  →  raw_cy = pg_h*(1-x_pct)
        #   rx     = W*(1-y_pct)  →  raw_cx = pg_w*(1-y_pct)
        raw_cx = pg_w * (1.0 - y_pct)
        raw_cy = pg_h * (1.0 - x_pct)

    # Bottom-left corner, clamped within page bounds.
    x_pt = max(0.0, min(raw_cx - raw_sig_w / 2.0, pg_w - raw_sig_w))
    y_pt = max(0.0, min(raw_cy - raw_sig_h / 2.0, pg_h - raw_sig_h))

    overlay_buf = _io.BytesIO()
    c = rl_canvas.Canvas(overlay_buf, pagesize=(pg_w, pg_h))
    c.drawImage(ImageReader(draw_img), x_pt, y_pt,
                width=raw_sig_w, height=raw_sig_h, mask="auto")
    c.save()
    return overlay_buf


# ── POST sign ─────────────────────────────────────────────────────────────────

@router.post("/{page_id}/files/page/{file_id}/sign")
async def sign_pdf(request: Request, page_id: int, file_id: int, body: SignBody):
    """Stamp a drawn signature onto the specified page of a PDF (in-place)."""
    if guard := _demo_guard(request):
        return guard
    uid = request.session.get("user_id")
    if not uid:
        raise HTTPException(status_code=401)
    await _require_uploads_page(page_id, uid)

    row = await get_page_upload_owned(file_id, uid)
    if not row or row["mime_type"] != "application/pdf":
        raise HTTPException(status_code=400, detail="File must be a PDF")

    try:
        _, b64 = body.signature_data.split(",", 1)
        sig_bytes = base64.b64decode(b64)
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Invalid signature data") from exc

    # Build the list of placements (multi-placement preferred; single legacy fallback)
    placements = body.placements or [
        SignPlacement(x_pct=body.x_pct, y_pct=body.y_pct, page_num=body.page_num)
    ]
    if not placements:
        raise HTTPException(status_code=400, detail="No signature placements provided")

    try:
        import PIL.Image
        import pypdf

        sig_img = PIL.Image.open(io.BytesIO(sig_bytes)).convert("RGBA")
        reader  = pypdf.PdfReader(str(UPLOAD_DIR / row["filename"]))
        writer  = pypdf.PdfWriter()
        writer.append(reader)
        page_count = len(writer.pages)

        for pl in placements:
            pg_idx      = min(max(pl.page_num, 0), page_count - 1)
            target_page = writer.pages[pg_idx]

            # _stamp_one_page handles rotation, coordinate transform, PIL pre-rotate
            overlay_buf = _stamp_one_page(sig_img, target_page, pl.x_pct, pl.y_pct)

            overlay_reader = pypdf.PdfReader(overlay_buf)
            target_page.merge_page(overlay_reader.pages[0])

        out_buf = io.BytesIO()
        writer.write(out_buf)
        data = out_buf.getvalue()
    except (HTTPException, ValueError):
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Signature stamping failed: {exc}") from exc

    # Save unsigned backup before first stamp so the user can undo
    bak_path = UPLOAD_DIR / (row["filename"] + ".bak")
    if not bak_path.exists():
        shutil.copy2(str(UPLOAD_DIR / row["filename"]), str(bak_path))

    (UPLOAD_DIR / row["filename"]).write_bytes(data)
    await update_page_upload_size(file_id, uid, len(data))
    return JSONResponse({"ok": True, "size": len(data), "has_backup": True, "stamps": len(placements)})


# ── GET /{pid}/files/page/{id}/wopi-token (Phase 6 — Feature A) ───────────────────

@router.get("/{page_id}/files/page/{file_id}/wopi-token")
async def wopi_token(
    request: Request, page_id: int, file_id: int,
):
    """Issue a short-lived WOPI access token for Collabora Online.

    Returns a JSON payload the JS uses to open the Collabora iframe:
      { token, wopi_src, editor_url, mime_type, original_name }

    If BW_COLLABORA_URL is not set, returns { collabora_disabled: true }
    so the JS can show a helpful message instead of a cryptic error.
    """
    # Import lazily to avoid circular import at module parse time
    from routers.wopi import (
        issue_token, get_editor_url, WOPI_MIMES,
        _COLLABORA_URL, _WOPI_BASE_URL,
    )

    uid = request.session.get("user_id")
    if not uid:
        raise HTTPException(status_code=401)

    if not _COLLABORA_URL:
        return JSONResponse({"collabora_disabled": True})

    await _require_uploads_page(page_id, uid)
    row = await get_page_upload_owned(file_id, uid)
    if not row:
        raise HTTPException(status_code=404, detail="File not found")

    mime = row.get("mime_type", "")
    if mime not in WOPI_MIMES:
        raise HTTPException(
            status_code=422,
            detail=f"MIME type '{mime}' is not supported by Collabora",
        )

    if not _WOPI_BASE_URL:
        raise HTTPException(
            status_code=503,
            detail="BW_WOPI_BASE_URL is not configured — Collabora cannot reach BookWorm",
        )

    token    = issue_token(file_id, uid, page_id)
    wopi_src = f"{_WOPI_BASE_URL}/wopi/files/{file_id}"

    try:
        editor_url = await get_editor_url(mime, wopi_src, token)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=503, detail=f"Could not build editor URL: {exc}") from exc

    return JSONResponse({
        "token":         token,
        "wopi_src":      wopi_src,
        "editor_url":    editor_url,
        "mime_type":     mime,
        "original_name": row.get("original_name", ""),
    })
